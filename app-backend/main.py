import asyncio
import base64
import io
import json
import os
import re
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import Optional

import httpx
import pypdf
from auth import (
    create_access_token,
    get_current_coach_email,
    get_current_user_email,
    get_optional_user_email,
    get_password_hash,
    verify_password,
)
from bson import ObjectId
from bson.errors import InvalidId
from docx import Document
from dotenv import load_dotenv
from fastapi import (
    APIRouter,
    BackgroundTasks,
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Query,
    UploadFile,
    status,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import OAuth2PasswordRequestForm
from models import (
    ChatRequest,
    Conversation,
    ConversationRename,
    Message,
    MessagePayload,
    PantryItem,
    PreferencesUpdate,
    ProfileUpdate,
    RecipeCreate,
    RosterLink,
    Token,
    UserCreate,
    UserInDB,
)
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from utils import extract_text_from_file, trim_conversation_history

load_dotenv()

# Global database client
db_client = None
db = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Connect to MongoDB on startup
    global db_client, db
    mongodb_uri = os.getenv("MONGODB_URI", "mongodb://localhost:27017/rizzbo")
    db_client = AsyncIOMotorClient(mongodb_uri)
    db = db_client.get_database()
    print("✅ Connected to MongoDB")

    yield

    # Clean up on shutdown
    db_client.close()
    print("❌ Disconnected from MongoDB")


app = FastAPI(title="RizzBo App Backend", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Conversation-Id"],
)


@app.get("/ping")
async def ping():
    return {"status": "ok", "service": "app-backend"}


@app.post("/register", status_code=status.HTTP_201_CREATED)
async def register_user(user: UserCreate):
    existing_user = await db.users.find_one({"email": user.email})
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="Email already registered"
        )

    hashed_pwd = get_password_hash(user.password)

    user_doc = {
        "email": user.email,
        "name": user.name.strip(),
        "hashed_password": hashed_pwd,
        "role": user.role,
        "plan": "Free plan",
        "target_weight": None,
        "activity_multiplier": 1.55,
        "preferences": {
            "theme": "system",
            "measurement_system": "metric",
            "dietary_preference": "none",
            "workout_reminders": True,
        },
        "created_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc),
    }

    await db.users.insert_one(user_doc)
    return {"message": "User successfully created"}


@app.post("/login", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user_doc = await db.users.find_one({"email": form_data.username})
    if not user_doc or not verify_password(
        form_data.password, user_doc["hashed_password"]
    ):
        raise HTTPException(status_code=400, detail="Incorrect email or password")

    token_payload = {
        "sub": user_doc["email"],
        "name": user_doc.get("name", "User"),
        "role": user_doc.get("role", "athlete"),
        "plan": user_doc.get("plan", "Free plan"),
    }

    access_token = create_access_token(data=token_payload)
    return {"access_token": access_token, "token_type": "bearer"}


@app.post("/upload/{conversation_id}")
async def upload_document(
    conversation_id: str,
    file: UploadFile = File(...),
    email: Optional[str] = Depends(get_optional_user_email),
):
    if not email:
        raise HTTPException(
            status_code=401, detail="Must be logged in to upload documents"
        )

    # Find User
    user = await db.users.find_one({"email": email})
    user_id = str(user["_id"]) if user else None

    # 1. Extract the text
    content = await file.read()
    extracted_text = await extract_text_from_file(content, file.filename)
    extracted_text = extracted_text[:15000]

    # 2. Prepare the Hidden Context Message (for the AI)
    sys_msg = {
        "role": "system",
        "content": f"[USER UPLOADED DOCUMENT: {file.filename}]\n\n{extracted_text}",
        "timestamp": datetime.now(timezone.utc),
    }

    # 3. Prepare the Visible Receipt Message (for the UI)
    ui_msg = {
        "role": "assistant",
        "content": f"📄 **Document Processed:** `{file.filename}`\n\nI have added this to my context. You can now ask me questions about it.",
        "timestamp": datetime.now(timezone.utc),
    }

    is_valid_id = ObjectId.is_valid(conversation_id)

    if not is_valid_id:
        new_db_id = ObjectId()
        new_conv = {
            "_id": new_db_id,
            "user_id": user_id,
            "title": f"{file.filename[:20]}",
            "messages": [sys_msg, ui_msg],
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc),
        }
        await db.conversations.insert_one(new_conv)

        return {
            "message": "Document processed and chat created",
            "filename": file.filename,
            "conversation_id": str(new_db_id),
        }
    else:
        result = await db.conversations.update_one(
            {"_id": ObjectId(conversation_id)},
            {
                "$push": {"messages": {"$each": [sys_msg, ui_msg]}},
                "$set": {"updated_at": datetime.now(timezone.utc)},
            },
        )

        if result.modified_count == 0:
            raise HTTPException(status_code=404, detail="Conversation not found")

        return {
            "message": "Document processed",
            "filename": file.filename,
            "conversation_id": conversation_id,
        }


async def generate_smart_title(first_prompt: str, ai_backend_url: str) -> str:
    """Generates a smart title using the AI and strictly enforces Sentence case."""
    system_instruction = (
        "You are a precise title generator. Create a concise, professional title (3-5 words) "
        "summarizing the user request. Do not include quotes, punctuation, or filler text."
    )
    title_prompt = f"{system_instruction}\n\nUser Request: {first_prompt}"

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{ai_backend_url}/chat",
                json={
                    "messages": [{"role": "user", "content": title_prompt}],
                    "model": "gemini-3.1-flash-lite",
                },
            )
            response.raise_for_status()

            # /chat returns an SSE stream — parse each "data: <json>" line properly
            # so that non-ASCII chars (đ/š/č/ć/ž) are decoded correctly instead of
            # staying as \uXXXX escape sequences.
            parts = []
            for line in response.text.splitlines():
                line = line.strip()
                if not line.startswith("data: "):
                    continue
                try:
                    chunk = json.loads(line[6:])
                    if isinstance(chunk, str):
                        parts.append(chunk)
                except Exception:
                    pass

            raw_title = "".join(parts).strip().replace('"', "").replace("'", "")
            return raw_title[:40].capitalize() or "New chat"

    except Exception:
        # Cast to str to prevent .split() crash on NoneTypes
        return " ".join(str(first_prompt).split()[:4]).capitalize() or "New chat"


async def extract_metrics_background(
    user_id: str, conversation_id: str, prompt: str, ai_backend_url: str
):
    extraction_prompt = (
        "You are an expert sports data extraction AI. Read the user message and extract important data into a strict JSON array.\n"
        "Categories to look for:\n"
        "1. 'body_stats': weight, height, body fat % (e.g., metric_name: 'weight').\n"
        "2. 'pr': personal records. Use metric_name for the exercise (e.g., 'deadlift_1rm'). Add {'sport': 'powerlifting'} to meta_data if known.\n"
        "3. 'goal': user's goals. value is the goal text. Add {'deadline': 'YYYY-MM-DD'} to meta_data if mentioned.\n"
        "4. 'training_data': frequency, duration, intensity, RPE (Rate of Perceived Exertion) (e.g., metric_name: 'session_duration', 'rpe').\n"
        "5. 'diet': calorie intake, macros, hydration (e.g., metric_name: 'daily_calories').\n"
        "6. 'recovery': sleep duration, sleep quality, soreness, fatigue levels (e.g., metric_name: 'sleep_hours', 'soreness_level').\n"
        "7. 'health': injuries, pain levels, resting heart rate, HRV (e.g., metric_name: 'knee_pain').\n\n"
        "Output ONLY valid JSON. If no metrics are found, output an empty array []. Example:\n"
        "[\n"
        '  {"category": "body_stats", "metric_name": "weight", "value": 82, "unit": "kg"},\n'
        '  {"category": "recovery", "metric_name": "sleep_hours", "value": 6, "unit": "hours", "meta_data": {"quality": "poor"}},\n'
        '  {"category": "training_data", "metric_name": "rpe", "value": 8, "unit": "", "meta_data": {"exercise": "sprints"}}\n'
        "]\n\n"
        f"User message: {prompt}"
    )

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                f"{ai_backend_url}/chat",
                json={
                    "messages": [{"role": "user", "content": extraction_prompt}],
                    "model": "gemini-3.1-flash-lite",
                },
            )
            response.raise_for_status()

            raw_json = ""
            for line in response.text.splitlines():
                if line.startswith("data: "):
                    try:
                        clean_chunk = json.loads(line.replace("data: ", ""))
                        raw_json += clean_chunk
                    except json.JSONDecodeError:
                        pass

            start_idx = raw_json.find("[")
            end_idx = raw_json.rfind("]")

            if start_idx != -1 and end_idx != -1:
                clean_json_string = raw_json[start_idx : end_idx + 1]

                if clean_json_string != "[]":
                    metrics = json.loads(clean_json_string)
                    for metric in metrics:
                        metric["user_id"] = user_id
                        metric["source_chat_id"] = conversation_id
                        metric["date"] = datetime.now(
                            timezone.utc
                        )  # Updated to timezone-aware UTC
                        if "meta_data" not in metric:
                            metric["meta_data"] = {}

                        await db.metrics.insert_one(metric)
                        print(
                            f"✅ Extracted: {metric['metric_name']} ({metric['value']})"
                        )
            else:
                print("No JSON array found in the AI response.")

    except Exception as e:
        print(f"⚠️ Background metric extraction failed: {repr(e)}")


async def update_long_term_memory(conversation_id: str, ai_backend_url: str):
    """Background task to update the conversation's rolling memory summary."""
    try:
        conv = await db.conversations.find_one({"_id": ObjectId(conversation_id)})
        if not conv:
            return

        messages = conv.get("messages", [])
        current_summary = conv.get("context_summary", "")

        # Strip datetime objects before JSON serialization
        recent_chunk = messages[-6:]
        clean_chunk = [
            {"role": m.get("role"), "content": m.get("content")} for m in recent_chunk
        ]

        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(
                f"{ai_backend_url}/summarize-memory",
                json={"messages": clean_chunk, "current_summary": current_summary},
            )
            resp.raise_for_status()
            new_summary = resp.json().get("summary", "")

        if new_summary:
            await db.conversations.update_one(
                {"_id": ObjectId(conversation_id)},
                {"$set": {"context_summary": new_summary}},
            )
            print(f"✅ Memory updated for conversation {conversation_id}")

    except Exception as e:
        print(f"⚠️ Memory Summarization Failed: {e}")


@app.post("/extract-text")
async def extract_text_only(
    file: UploadFile = File(...),
    email: Optional[str] = Depends(get_optional_user_email),
):
    """
    Isolated endpoint for extracting text from staging files.
    Does NOT interact with the database or LLM.
    """
    if not email:
        raise HTTPException(
            status_code=401, detail="Must be logged in to process files."
        )

    try:
        content = await file.read()
        filename = file.filename.lower()
        extracted_text = ""

        if filename.endswith(".pdf"):
            # Load the binary stream into the pypdf reader
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"

        elif filename.endswith((".txt", ".md", ".csv")):
            # Standard UTF-8 decoding for plaintext files
            extracted_text = content.decode("utf-8")

        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload PDF, TXT, MD, or CSV.",
            )

        # SECURITY & PERFORMANCE GUARDRAIL:
        # Hard cap the extracted text at ~15,000 characters to prevent
        # massive documents from exceeding the LLM context window or
        # slowing down the inference stream.
        extracted_text = extracted_text[:15000]

        return {"filename": file.filename, "text": extracted_text}

    except UnicodeDecodeError:
        raise HTTPException(
            status_code=400, detail="File encoding not supported. Please use UTF-8."
        )
    except Exception as e:
        print(f"Extraction error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse document.")


@app.post("/chat")
async def chat(
    request: ChatRequest,
    email: Optional[str] = Depends(get_optional_user_email),
):
    user_id = None
    role = "athlete"
    if email:
        user = await db.users.find_one({"email": email})
        if user:
            user_id = str(user["_id"])
            role = user.get("role", "athlete")
            measurement_system = user.get("preferences", {}).get("measurement_system", "metric")
    else:
        measurement_system = "metric"


    is_valid_id = (
        ObjectId.is_valid(request.conversation_id) if request.conversation_id else False
    )
    is_new_conversation = not request.conversation_id or not is_valid_id

    if is_new_conversation:
        db_assigned_id = ObjectId()
        current_conv_id = str(db_assigned_id)
        if user_id:
            shell_conv = {
                "_id": db_assigned_id,
                "user_id": user_id,
                "title": request.prompt[:30] + "...",
                "messages": [],
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }
            await db.conversations.insert_one(shell_conv)
    else:
        current_conv_id = request.conversation_id
        db_assigned_id = ObjectId(current_conv_id)

    async def generate():
        user_msg = {
            "role": "user",
            "content": request.prompt,
            "timestamp": datetime.now(timezone.utc),
        }
        ai_content = ""
        ai_backend_url = os.getenv("AI_BACKEND_URL", "http://127.0.0.1:8000")

        title_task = None
        if is_new_conversation:
            title_task = asyncio.create_task(
                generate_smart_title(request.prompt, ai_backend_url)
            )

        # --- 1. GATHER DYNAMIC CONTEXT (Roster & Memory) ---
        roster_context = ""
        if role == "coach" and user_id:
            # (Your existing roster fetching logic remains exactly the same here)
            links = await db.roster_links.find(
                {"coach_id": user_id, "status": "active"}
            ).to_list(length=100)
            if links:
                athlete_ids = [link["athlete_id"] for link in links]
                pipeline = [
                    {
                        "$match": {
                            "user_id": {"$in": athlete_ids},
                            "metric_name": {"$exists": True, "$ne": None},
                        }
                    },
                    {"$sort": {"date": -1}},
                    {
                        "$group": {
                            "_id": {
                                "user_id": "$user_id",
                                "metric_name": "$metric_name",
                            },
                            "latest_value": {"$first": "$value"},
                            "unit": {"$first": "$unit"},
                        }
                    },
                ]
                metrics_cursor = db.metrics.aggregate(pipeline)
                team_metrics = await metrics_cursor.to_list(length=500)
                valid_athlete_obj_ids = [
                    ObjectId(aid) for aid in athlete_ids if ObjectId.is_valid(aid)
                ]
                athletes = await db.users.find(
                    {"_id": {"$in": valid_athlete_obj_ids}}, {"name": 1}
                ).to_list(length=100)
                athlete_map = {
                    str(a["_id"]): a.get("name", "Unknown") for a in athletes
                }

                formatted_stats = {}
                for m in team_metrics:
                    group_id = m.get("_id", {})
                    uid = group_id.get("user_id")
                    metric = group_id.get("metric_name")
                    if not uid or not metric:
                        continue
                    a_name = athlete_map.get(uid, "Unknown")
                    if a_name not in formatted_stats:
                        formatted_stats[a_name] = []
                    formatted_stats[a_name].append(
                        f"{metric}: {m.get('latest_value', '')}{m.get('unit', '')}"
                    )

                if formatted_stats:
                    roster_context = "[ROSTER CONTEXT - LATEST TEAM DATA]:\n"
                    for a_name, stats in formatted_stats.items():
                        roster_context += f"- {a_name}: {', '.join(stats)}\n"
                    roster_context += "[END ROSTER CONTEXT]\n\n"

        # --- 2. BUILD STRUCTURED MESSAGE ARRAY ---
        payload_messages = []
        memory_block = ""

        if not is_new_conversation:
            try:
                conv = await db.conversations.find_one({"_id": db_assigned_id})
                if conv and "messages" in conv:
                    context_summary = conv.get("context_summary", "")
                    if context_summary:
                        memory_block = (
                            f"[LONG-TERM MEMORY SUMMARY]:\n{context_summary}\n\n"
                        )

                    # Separate system/doc messages from chat
                    doc_messages = [
                        m for m in conv["messages"] if m.get("role") == "system"
                    ]
                    chat_messages = [
                        m for m in conv["messages"] if m.get("role") != "system"
                    ]

                    # Trim chat history (you should update your trim_conversation_history util
                    # to return a list of dicts instead of formatting strings)
                    trimmed_history = trim_conversation_history(
                        chat_messages, max_words=2500
                    )

                    # Inject pinned documents as a system-like preamble to the first message
                    doc_preamble = ""
                    for m in doc_messages:
                        doc_preamble += (
                            f"PINNED DOCUMENT CONTEXT: {m.get('content', '')}\n\n"
                        )

                    # SANITIZER: Ensure strictly alternating roles (user -> model)
                    last_role = None
                    for m in trimmed_history:
                        m_role = "user" if m.get("role") == "user" else "model"

                        # Skip consecutive messages from the same role to prevent Google 400 errors
                        if m_role == last_role:
                            continue

                        # If this is the very first message in the payload, inject the doc_preamble
                        content = m.get("content", "")
                        if last_role is None and doc_preamble:
                            content = f"{doc_preamble}{content}"

                        payload_messages.append({"role": m_role, "content": content})
                        last_role = m_role

            except Exception as e:
                print(f"Error loading history: {e}")

        # --- 3. CONSTRUCT THE LATEST PROMPT WITH INJECTED CONTEXT ---
        # We wrap the Roster & Memory specifically around the final user prompt.
        # This acts as dynamic context (RAG) without polluting the history array.
        unit_hint = (
            "Use metric units (kg, cm, km, ml)."
            if measurement_system == "metric"
            else "Use imperial units (lbs, inches, miles, fl oz)."
        )
        measurement_block = f"[USER PREFERENCE] {unit_hint}\n\n"

        final_prompt_content = (
            f"{measurement_block}{memory_block}{roster_context}USER PROMPT: {request.prompt}"
        )

        # Ensure the final appended message doesn't break the user->model->user alternation
        if payload_messages and payload_messages[-1]["role"] == "user":
            # Edge case: Last saved message was a user (e.g. generation failed previously)
            # Replace it with the new prompt
            payload_messages[-1] = {"role": "user", "content": final_prompt_content}
        else:
            payload_messages.append({"role": "user", "content": final_prompt_content})

        # --- 4. STREAM TO AI-BACKEND ---
        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                async with client.stream(
                    "POST",
                    f"{ai_backend_url}/chat",
                    # Send the structured array instead of a single string
                    json={
                        "messages": payload_messages,
                        "model": request.model,
                        "enable_search": request.enable_search,
                        # Clean original prompt — avoids context injections polluting the search query
                        "search_query": request.prompt if request.enable_search else None,
                    }
                ) as response:
                    response.raise_for_status()
                    async for line in response.aiter_lines():
                        if line.startswith("data: "):
                            try:
                                clean_text = json.loads(line.replace("data: ", ""))
                                ai_content += clean_text
                            except json.JSONDecodeError:
                                pass
                        yield line + "\n"
        except (httpx.HTTPStatusError, httpx.RequestError) as e:
            print(f"❌ AI Backend connection failure: {e}")
            mock_words = ["I", " couldn't", " reach", " the", " AI..."]
            for word in mock_words:
                ai_content += word
                yield f"data: {json.dumps(word)}\n\n"
                await asyncio.sleep(0.1)

        # --- 5. SAVE FINALIZED MESSAGES TO DB ---
        # Strip the auto-search marker before persisting — it's a frontend signal, not content.
        ai_content = ai_content.replace("[NEEDS_WEB_SEARCH]", "").strip()

        # is_retry=True means this was an auto-retry triggered by [NEEDS_WEB_SEARCH].
        # We skip DB save to avoid duplicate user messages; the initial response is already saved.
        if user_id and not request.is_retry:
            ai_msg = {
                "role": "assistant",
                "content": ai_content,
                "timestamp": datetime.utcnow(),
            }
            await db.conversations.update_one(
                {"_id": db_assigned_id},
                {
                    "$push": {"messages": {"$each": [user_msg, ai_msg]}},
                    "$set": {"updated_at": datetime.utcnow()},
                },
            )

            # --- CALCULATE MESSAGE COUNT FOR LONG-TERM MEMORY TRIGGER ---
            # If it's a new conversation, we just added 2 messages.
            # Otherwise, we fetch the length of the existing messages from the 'conv' object and add 2.
            try:
                current_msg_count = (
                    len(conv.get("messages", [])) + 2 if not is_new_conversation else 2
                )
            except NameError:
                # Fallback just in case 'conv' wasn't loaded
                current_msg_count = 2

            if title_task:
                try:
                    generated_title = await title_task
                    await db.conversations.update_one(
                        {"_id": db_assigned_id}, {"$set": {"title": generated_title}}
                    )
                    yield f"data: {json.dumps({'_type': 'title_update', 'title': generated_title})}\n\n"
                except Exception as e:
                    pass

            # --- FIRE-AND-FORGET BACKGROUND TASKS (Bypassing FastAPI Streaming limitations) ---

            # 1. Trigger the RAG Summarization every 6 messages
            if current_msg_count > 0 and current_msg_count % 6 == 0:
                asyncio.create_task(
                    update_long_term_memory(str(db_assigned_id), ai_backend_url)
                )

            # 2. Trigger the Metric Extraction on every single message
            asyncio.create_task(
                extract_metrics_background(
                    user_id, current_conv_id, request.prompt, ai_backend_url
                )
            )

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"X-Conversation-Id": current_conv_id},
    )


@app.get("/conversations")
async def get_all_conversations(email: str = Depends(get_current_user_email)):
    # 1. Find the user
    user = await db.users.find_one({"email": email})

    # SAFEGUARD: Check if the user actually exists before trying to access user["_id"]
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="User not found. Please log in again.",
        )

    # 2. Fetch their conversations (excluding the heavy messages array to save bandwidth)
    cursor = db.conversations.find({"user_id": str(user["_id"])}, {"messages": 0}).sort(
        "updated_at", -1
    )
    conversations = await cursor.to_list(length=100)

    # 3. Convert MongoDB ObjectIds to strings for JSON
    for conv in conversations:
        conv["_id"] = str(conv["_id"])

    return conversations


@app.get("/conversations/{conversation_id}")
async def get_single_conversation(
    conversation_id: str, email: str = Depends(get_current_user_email)
):
    user = await db.users.find_one({"email": email})
    conv = await db.conversations.find_one(
        {"_id": ObjectId(conversation_id), "user_id": str(user["_id"])}
    )

    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    conv["_id"] = str(conv["_id"])
    return conv


@app.patch("/conversations/{conversation_id}")
async def rename_conversation(
    conversation_id: str,
    payload: ConversationRename,
    email: str = Depends(get_current_user_email),
):
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    user_id = str(user["_id"])

    if not ObjectId.is_valid(conversation_id):
        raise HTTPException(status_code=400, detail="Invalid conversation ID format")

    # Update the title, but ONLY if the user_id matches the owner!
    result = await db.conversations.update_one(
        {"_id": ObjectId(conversation_id), "user_id": user_id},
        {
            "$set": {
                "title": payload.title.strip(),
                "updated_at": datetime.now(timezone.utc),
            }
        },
    )

    if result.matched_count == 0:
        raise HTTPException(
            status_code=404,
            detail="Conversation not found or you do not have permission to rename it.",
        )

    return {
        "message": "Conversation renamed successfully",
        "title": payload.title.strip(),
    }


@app.delete("/conversations/{conversation_id}")
async def delete_conversation(
    conversation_id: str, email: str = Depends(get_current_user_email)
):
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    user_id = str(user["_id"])

    if not ObjectId.is_valid(conversation_id):
        raise HTTPException(status_code=400, detail="Invalid conversation ID format")

    # 1. First, delete the conversation
    result = await db.conversations.delete_one(
        {"_id": ObjectId(conversation_id), "user_id": user_id}
    )

    if result.deleted_count == 0:
        raise HTTPException(
            status_code=404,
            detail="Conversation not found or you do not have permission to delete it.",
        )

    # 2. Next, CASCADE DELETE: Remove all metrics that were extracted during this chat!
    # (This prevents phantom data from showing up on the Profile page)
    await db.metrics.delete_many({"source_chat_id": ObjectId(conversation_id)})

    return {"message": "Conversation and associated metrics deleted successfully"}


@app.get("/search")
async def search_conversations(q: str, email: str = Depends(get_current_user_email)):
    user = await db.users.find_one({"email": email})

    # Search for the keyword inside the messages content using regex (case-insensitive)
    query = {
        "user_id": str(user["_id"]),
        "messages.content": {"$regex": q, "$options": "i"},
    }

    cursor = db.conversations.find(query, {"messages": 0}).sort("updated_at", -1)
    results = await cursor.to_list(length=50)

    for res in results:
        res["_id"] = str(res["_id"])

    return results


@app.get("/export/{conversation_id}")
async def export_conversation(
    conversation_id: str,
    format: str = Query("docx", regex="^(docx|pdf)$"),  # Restricts input to docx or pdf
    email: str = Depends(get_current_user_email),
):
    user = await db.users.find_one({"email": email})
    conv = await db.conversations.find_one(
        {"_id": ObjectId(conversation_id), "user_id": str(user["_id"])}
    )

    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    file_stream = io.BytesIO()

    if format == "pdf":
        # --- PDF GENERATION ---
        p = canvas.Canvas(file_stream, pagesize=letter)
        p.drawString(100, 750, f"Title: {conv.get('title', 'Chat Export')}")
        y = 700
        for msg in conv.get("messages", []):
            if msg["role"] == "system":
                continue
            text = f"{msg['role'].capitalize()}: {msg['content'][:100]}"  # Truncate for simplicity
            p.drawString(100, y, text)
            y -= 20
        p.save()
        filename = "RizzBo_Chat.pdf"
        media_type = "application/pdf"
    else:
        # --- DOCX GENERATION ---
        document = Document()
        document.add_heading(conv.get("title", "RizzBo Chat Export"), 0)
        for msg in conv.get("messages", []):
            if msg["role"] == "system":
                continue
            p = document.add_paragraph()
            p.add_run(f"{msg['role'].capitalize()}: ").bold = True
            p.add_run(msg["content"])
        document.save(file_stream)
        filename = "RizzBo_Chat.docx"
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    file_stream.seek(0)
    return StreamingResponse(
        file_stream,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.get("/profile/summary")
async def get_profile_summary(email: str = Depends(get_current_user_email)):
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    pipeline = [
        {"$match": {"user_id": str(user["_id"])}},
        {"$sort": {"date": -1}},
        {
            "$group": {
                "_id": "$metric_name",
                "latest_value": {"$first": "$value"},
                "unit": {"$first": "$unit"},
                "category": {"$first": "$category"},
                "meta_data": {"$first": "$meta_data"},
                "date": {"$first": "$date"},
            }
        },
    ]

    cursor = db.metrics.aggregate(pipeline)
    results = await cursor.to_list(length=100)

    summary = {}
    for r in results:
        summary[r["_id"]] = {
            "value": r["latest_value"],
            "unit": r.get("unit", ""),
            "category": r.get("category", "general"),
            "meta_data": r.get("meta_data", {}),
            "date": r["date"].isoformat(),
        }
    return summary


@app.get("/profile/metrics/{metric_name}")
async def get_profile_metrics(
    metric_name: str, email: str = Depends(get_current_user_email)
):
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    # Sort ascending so the frontend charts render chronologically
    cursor = db.metrics.find(
        {"user_id": str(user["_id"]), "metric_name": metric_name}, {"_id": 0}
    ).sort("date", 1)

    metrics = await cursor.to_list(length=100)
    # Convert dates to ISO strings for JSON serialization
    for m in metrics:
        if isinstance(m.get("date"), datetime):
            m["date"] = m["date"].isoformat()

    return metrics


@app.post("/roster/invite")
async def invite_athlete(
    payload: dict, coach_email: str = Depends(get_current_coach_email)
):
    athlete_email = payload.get("email")
    if not athlete_email:
        raise HTTPException(status_code=400, detail="Athlete email is required")

    coach = await db.users.find_one({"email": coach_email})
    athlete = await db.users.find_one({"email": athlete_email, "role": "athlete"})

    if not athlete:
        raise HTTPException(
            status_code=404,
            detail="Athlete not found, or user is not registered as an athlete.",
        )

    # Check if they are already linked
    existing_link = await db.roster_links.find_one(
        {"coach_id": str(coach["_id"]), "athlete_id": str(athlete["_id"])}
    )

    if existing_link:
        if existing_link.get("status") == "rejected":
            # Allow the coach to re-send the invite by resetting it to pending
            await db.roster_links.update_one(
                {"_id": existing_link["_id"]}, {"$set": {"status": "pending"}}
            )
            return {"message": "Invite re-sent! Waiting for athlete approval."}
        else:
            raise HTTPException(
                status_code=400,
                detail="Athlete has already been invited or is on your roster.",
            )

    new_link = RosterLink(
        coach_id=str(coach["_id"]), athlete_id=str(athlete["_id"]), status="pending"
    )

    await db.roster_links.insert_one(new_link.model_dump())
    return {"message": "Invite sent! Waiting for athlete approval."}


@app.get("/roster/athletes")
async def get_roster(coach_email: str = Depends(get_current_coach_email)):
    coach = await db.users.find_one({"email": coach_email})

    # 1. Find all links belonging to this coach
    cursor = db.roster_links.find({"coach_id": str(coach["_id"])})
    links = await cursor.to_list(length=100)

    if not links:
        return []

    # 2. Extract athlete IDs
    from bson.objectid import ObjectId

    athlete_ids = [ObjectId(link["athlete_id"]) for link in links]

    # 3. Fetch the athlete's profiles (Excluding passwords!)
    athletes_cursor = db.users.find(
        {"_id": {"$in": athlete_ids}}, {"hashed_password": 0}
    )
    athletes = await athletes_cursor.to_list(length=100)

    # 4. Format for the frontend
    roster = []
    for a in athletes:
        status = next(
            (l["status"] for l in links if l["athlete_id"] == str(a["_id"])), "unknown"
        )
        roster.append(
            {
                "id": str(a["_id"]),
                "name": a.get("name", "Unknown Athlete"),
                "email": a.get("email"),
                "status": status,
            }
        )

    return roster


# --- ATHLETE CONSENT ENDPOINTS ---


@app.get("/athlete/invites")
async def get_pending_invites(email: str = Depends(get_current_user_email)):
    athlete = await db.users.find_one({"email": email})
    athlete_id = str(athlete["_id"])

    # Find all pending links for this athlete
    cursor = db.roster_links.find({"athlete_id": athlete_id, "status": "pending"})
    links = await cursor.to_list(length=100)

    if not links:
        return []

    # Extract coach IDs to get their names
    from bson.objectid import ObjectId

    coach_ids = [ObjectId(link["coach_id"]) for link in links]

    coaches_cursor = db.users.find({"_id": {"$in": coach_ids}}, {"name": 1, "email": 1})
    coaches = await coaches_cursor.to_list(length=100)

    results = []
    for coach in coaches:
        results.append(
            {
                "coach_id": str(coach["_id"]),
                "name": coach.get("name", "Unknown Coach"),
                "email": coach.get("email"),
            }
        )

    return results


@app.get("/athlete/coaches")
async def get_active_coaches(email: str = Depends(get_current_user_email)):
    athlete = await db.users.find_one({"email": email})

    cursor = db.roster_links.find(
        {"athlete_id": str(athlete["_id"]), "status": "active"}
    )
    links = await cursor.to_list(length=100)

    if not links:
        return []

    from bson.objectid import ObjectId

    coach_ids = [ObjectId(link["coach_id"]) for link in links]

    coaches_cursor = db.users.find({"_id": {"$in": coach_ids}}, {"name": 1})
    coaches = await coaches_cursor.to_list(length=100)

    results = []
    for coach in coaches:
        results.append(
            {"id": str(coach["_id"]), "name": coach.get("name", "Unknown Coach")}
        )

    return results


@app.post("/athlete/invites/{coach_id}/respond")
async def respond_to_invite(
    coach_id: str, payload: dict, email: str = Depends(get_current_user_email)
):
    athlete = await db.users.find_one({"email": email})
    athlete_id = str(athlete["_id"])
    action = payload.get("action")

    if action not in ["accept", "reject"]:
        raise HTTPException(
            status_code=400, detail="Invalid action. Use 'accept' or 'reject'."
        )

    link = await db.roster_links.find_one(
        {"coach_id": coach_id, "athlete_id": athlete_id, "status": "pending"}
    )

    if not link:
        raise HTTPException(status_code=404, detail="Pending invite not found.")

    if action == "accept":
        await db.roster_links.update_one(
            {"_id": link["_id"]}, {"$set": {"status": "active"}}
        )
        return {"message": "Invite accepted. Coach now has access to your data."}
    else:
        # FIX: Keep the record, but mark it as rejected
        await db.roster_links.update_one(
            {"_id": link["_id"]}, {"$set": {"status": "rejected"}}
        )
        return {"message": "Invite rejected."}


@app.get("/coach/metrics/summary")
async def get_coach_metrics_summary(
    coach_email: str = Depends(get_current_coach_email),
):
    coach = await db.users.find_one({"email": coach_email})
    if not coach:
        raise HTTPException(
            status_code=401, detail="Coach account no longer exists in database."
        )

    links_cursor = db.roster_links.find(
        {"coach_id": str(coach["_id"]), "status": "active"}
    )
    links = await links_cursor.to_list(length=100)

    if not links:
        return []

    athlete_ids = [link["athlete_id"] for link in links]

    pipeline = [
        {
            "$match": {
                "user_id": {"$in": athlete_ids},
                "metric_name": {
                    "$exists": True,
                    "$ne": None,
                },
            }
        },
        {"$sort": {"date": -1}},
        {
            "$group": {
                "_id": {"user_id": "$user_id", "metric_name": "$metric_name"},
                "latest_value": {"$first": "$value"},
                "unit": {"$first": "$unit"},
                "category": {"$first": "$category"},
                "meta_data": {"$first": "$meta_data"},
                "date": {"$first": "$date"},
            }
        },
    ]

    cursor = db.metrics.aggregate(pipeline)
    results = await cursor.to_list(length=500)

    valid_athlete_object_ids = []
    for aid in athlete_ids:
        if ObjectId.is_valid(aid):
            valid_athlete_object_ids.append(ObjectId(aid))

    athletes_cursor = db.users.find(
        {"_id": {"$in": valid_athlete_object_ids}}, {"name": 1}
    )
    athletes = await athletes_cursor.to_list(length=100)
    athlete_map = {str(a["_id"]): a.get("name", "Unknown Athlete") for a in athletes}

    summary_by_athlete = {}
    for r in results:
        group_id = r.get("_id", {})
        uid = group_id.get("user_id")
        metric = group_id.get("metric_name")

        if not uid or not metric:
            continue

        if uid not in summary_by_athlete:
            summary_by_athlete[uid] = {
                "athlete_id": uid,
                "name": athlete_map.get(uid, "Unknown Athlete"),
                "metrics": {},
            }

        raw_date = r.get("date")
        if isinstance(raw_date, datetime):
            safe_date_str = raw_date.isoformat()
        elif isinstance(raw_date, str):
            safe_date_str = raw_date
        else:
            safe_date_str = datetime.now(timezone.utc).isoformat()

        summary_by_athlete[uid]["metrics"][metric] = {
            "value": r.get("latest_value"),
            "unit": r.get("unit", ""),
            "category": r.get("category", "general"),
            "meta_data": r.get("meta_data", {}),
            "date": safe_date_str,
        }

    return list(summary_by_athlete.values())


@app.get("/coach/metrics/{athlete_id}")
async def get_coach_athlete_metrics(
    athlete_id: str,
    metric_name: str,
    coach_email: str = Depends(get_current_coach_email),
):
    coach = await db.users.find_one({"email": coach_email})

    link = await db.roster_links.find_one(
        {"coach_id": str(coach["_id"]), "athlete_id": athlete_id, "status": "active"}
    )

    if not link:
        raise HTTPException(
            status_code=403, detail="Access denied. This athlete is not on your roster."
        )

    # Fetch time-series data for the chart
    cursor = db.metrics.find(
        {"user_id": athlete_id, "metric_name": metric_name}, {"_id": 0}
    ).sort("date", 1)

    metrics = await cursor.to_list(length=100)

    for m in metrics:
        if isinstance(m.get("date"), datetime):
            m["date"] = m["date"].isoformat()

    return metrics


# --- SMART KITCHEN & PANTRY ENDPOINTS ---


@app.get("/kitchen/pantry")
async def get_pantry(email: str = Depends(get_current_user_email)):
    user = await db.users.find_one({"email": email})
    cursor = db.pantry.find({"user_id": str(user["_id"])})
    items = await cursor.to_list(length=200)

    for item in items:
        item["_id"] = str(item["_id"])
    return items


@app.post("/kitchen/pantry")
async def add_pantry_item(payload: dict, email: str = Depends(get_current_user_email)):
    user = await db.users.find_one({"email": email})

    item_name = payload.get("item_name")
    if not item_name:
        raise HTTPException(status_code=400, detail="item_name is required")

    item = PantryItem(
        user_id=str(user["_id"]),
        item_name=item_name,
        quantity=payload.get("quantity", ""),
    )

    # Upsert logic: if it exists, update it. If not, insert it.
    await db.pantry.update_one(
        {"user_id": str(user["_id"]), "item_name": item_name},
        {"$set": item.model_dump()},
        upsert=True,
    )
    return {"message": "Pantry updated"}


@app.delete("/kitchen/pantry/{item_id}")
async def delete_pantry_item(
    item_id: str, email: str = Depends(get_current_user_email)
):
    user = await db.users.find_one({"email": email})
    result = await db.pantry.delete_one(
        {"_id": ObjectId(item_id), "user_id": str(user["_id"])}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Item not found")
    return {"message": "Item removed"}


@app.post("/kitchen/generate")
async def generate_recipe(
    image: Optional[UploadFile] = File(None),
    text_ingredients: Optional[str] = Form(None),
    target_protein: Optional[float] = Form(None),
    target_carbs: Optional[float] = Form(None),
    email: str = Depends(get_current_user_email),
):
    try:
        # 1. Safe User Lookup
        user = await db.users.find_one({"email": email})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        user_id = str(user["_id"])

        # 2. Fetch Pantry Items
        pantry_cursor = db.pantry.find({"user_id": user_id})
        pantry_items = await pantry_cursor.to_list(length=200)
        pantry_list = ", ".join(
            [
                f"{item['item_name']} ({item.get('quantity', '')})"
                for item in pantry_items
            ]
        )

        # 3. Read Image safely
        images_b64 = []
        if image:
            image_bytes = await image.read()
            b64_string = base64.b64encode(image_bytes).decode("utf-8")
            images_b64.append(b64_string)

        # 4. Construct Prompt
        prompt = "You are an expert sports nutritionist AI. Generate a precise recipe based on the following.\n"
        if pantry_list:
            prompt += f"Always Available Pantry Ingredients: {pantry_list}\n"
        if text_ingredients:
            prompt += f"Specific Ingredients requested: {text_ingredients}\n"
        if image:
            prompt += "Also use the ingredients visible in the provided image.\n"

        prompt += "\nTARGET MACROS FOR THIS MEAL:\n"
        if target_protein:
            prompt += f"- Protein: ~{target_protein}g\n"
        if target_carbs:
            prompt += f"- Carbs: ~{target_carbs}g\n"

        prompt += (
            "\nIMPORTANT: Return ONLY a valid JSON object matching this exact structure. Do not include markdown formatting.\n"
            "{\n"
            '  "title": "Recipe Name",\n'
            '  "prep_time_minutes": 15,\n'
            '  "macros": {"protein": 50, "carbs": 60, "fats": 10, "calories": 550},\n'
            '  "ingredients": ["100g Chicken", "50g Rice"],\n'
            '  "instructions": ["Step 1...", "Step 2..."]\n'
            "}"
        )

        # 5. Call AI Backend
        ai_backend_url = os.getenv("AI_BACKEND_URL", "http://127.0.0.1:8000")
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{ai_backend_url}/generate-recipe",
                json={"prompt": prompt, "images": images_b64},
            )
            response.raise_for_status()

            # 6. Parse and Clean the LLM Output
            ai_data = response.json()

            # Assuming your AI microservice returns something like {"response": "{...}"}
            raw_text = (
                ai_data.get("response", "")
                if isinstance(ai_data, dict)
                else str(ai_data)
            )

            # Strip markdown code blocks (```json ... ```) just in case
            clean_json_str = re.sub(r"```json\s*|\s*```", "", raw_text).strip()

            # Validate it's actually JSON before sending to frontend
            parsed_recipe = json.loads(clean_json_str)
            return parsed_recipe

    except httpx.HTTPStatusError as e:
        print(f"AI Backend HTTP Error: {e.response.text}")
        raise HTTPException(
            status_code=502, detail="AI microservice failed to respond correctly."
        )
    except json.JSONDecodeError as e:
        print(f"LLM Output Parsing Error: {clean_json_str}")
        raise HTTPException(
            status_code=500, detail="AI returned malformed data. Try generating again."
        )
    except Exception as e:
        print(f"Internal Generation Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# --- USER SETTINGS ENDPOINTS ---
@app.get("/users/me")
async def get_current_user_profile(email: str = Depends(get_current_user_email)):
    user = await db.users.find_one({"email": email}, {"hashed_password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Convert ObjectId for JSON serialization
    user["_id"] = str(user["_id"])
    return user


@app.patch("/profile")
async def update_profile(
    payload: ProfileUpdate, email: str = Depends(get_current_user_email)
):
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    update_data = {}
    if payload.name is not None:
        update_data["name"] = payload.name.strip()
    if payload.target_weight is not None:
        update_data["target_weight"] = payload.target_weight
    if payload.activity_multiplier is not None:
        update_data["activity_multiplier"] = payload.activity_multiplier

    if update_data:
        update_data["updated_at"] = datetime.now(timezone.utc)
        await db.users.update_one({"_id": user["_id"]}, {"$set": update_data})

    return {"message": "Profile updated", "updated_fields": update_data}


@app.patch("/preferences")
async def update_preferences(
    payload: PreferencesUpdate, email: str = Depends(get_current_user_email)
):
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    update_data = {}
    # Use MongoDB dot notation to update specific nested fields safely
    if payload.theme is not None:
        update_data["preferences.theme"] = payload.theme
    if payload.measurement_system is not None:
        update_data["preferences.measurement_system"] = payload.measurement_system
    if payload.dietary_preference is not None:
        update_data["preferences.dietary_preference"] = payload.dietary_preference
    if payload.workout_reminders is not None:
        update_data["preferences.workout_reminders"] = payload.workout_reminders

    if update_data:
        update_data["updated_at"] = datetime.now(timezone.utc)
        await db.users.update_one({"_id": user["_id"]}, {"$set": update_data})

    return {"message": "Preferences updated", "updated_fields": update_data}


@app.post("/recipes", status_code=status.HTTP_201_CREATED)
async def save_recipe(
    recipe: RecipeCreate, email: str = Depends(get_current_user_email)
):
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    recipe_doc = recipe.model_dump()

    recipe_doc["user_id"] = str(user["_id"])
    recipe_doc["created_at"] = datetime.now(timezone.utc)

    result = await db.recipes.insert_one(recipe_doc)
    return {"message": "Recipe saved", "id": str(result.inserted_id)}


@app.get("/recipes")
async def get_saved_recipes(email: str = Depends(get_current_user_email)):
    user = await db.users.find_one({"email": email})
    cursor = db.recipes.find({"user_id": str(user["_id"])}).sort("created_at", -1)
    recipes = await cursor.to_list(length=100)

    for r in recipes:
        r["_id"] = str(r["_id"])
    return recipes


@app.delete("/recipes/{recipe_id}")
async def delete_recipe(recipe_id: str, email: str = Depends(get_current_user_email)):
    user = await db.users.find_one({"email": email})
    if not ObjectId.is_valid(recipe_id):
        raise HTTPException(status_code=400, detail="Invalid recipe ID")

    result = await db.recipes.delete_one(
        {"_id": ObjectId(recipe_id), "user_id": str(user["_id"])}
    )

    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Recipe not found")
    return {"message": "Recipe removed"}


class PDFGenerateRequest(BaseModel):
    content: str


# Add endpoint
@app.post("/generate-pdf")
async def generate_pdf(request: PDFGenerateRequest):
    file_stream = io.BytesIO()

    doc = SimpleDocTemplate(
        file_stream,
        pagesize=letter,
        rightMargin=60, leftMargin=60,
        topMargin=60, bottomMargin=60,
    )

    styles = getSampleStyleSheet()

    # --- Custom styles ---
    h1 = ParagraphStyle("H1", parent=styles["Normal"], fontSize=20, leading=26, spaceBefore=14, spaceAfter=6, textColor=colors.HexColor("#111111"), fontName="Helvetica-Bold")
    h2 = ParagraphStyle("H2", parent=styles["Normal"], fontSize=16, leading=22, spaceBefore=12, spaceAfter=4, textColor=colors.HexColor("#222222"), fontName="Helvetica-Bold")
    h3 = ParagraphStyle("H3", parent=styles["Normal"], fontSize=13, leading=18, spaceBefore=10, spaceAfter=3, textColor=colors.HexColor("#333333"), fontName="Helvetica-Bold")
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=16, spaceAfter=6, fontName="Helvetica")
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10, leading=16, leftIndent=20, firstLineIndent=0, spaceAfter=3, fontName="Helvetica")
    code_style = ParagraphStyle("Code", parent=styles["Normal"], fontSize=9, leading=14, leftIndent=20, fontName="Courier", backColor=colors.HexColor("#f4f4f4"), spaceAfter=6)

    TABLE_STYLE = TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0),  colors.HexColor("#2d2d2d")),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#ffffff"), colors.HexColor("#f2f2f2")]),
        ("GRID",         (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
    ])

    def md_inline(text: str) -> str:
        """Convert inline markdown (bold, italic, code, links) to ReportLab XML."""
        # Inline code  `code`  → <font> with courier
        text = re.sub(r"`([^`]+)`", r'<font name="Courier" size="9">\1</font>', text)
        # Bold+italic ***text***
        text = re.sub(r"\*\*\*(.*?)\*\*\*", r"<b><i>\1</i></b>", text)
        # Bold **text** or __text__
        text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"__(.*?)__", r"<b>\1</b>", text)
        # Italic *text* or _text_
        text = re.sub(r"(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text)
        text = re.sub(r"(?<!_)_(?!_)(.*?)(?<!_)_(?!_)", r"<i>\1</i>", text)
        # Links [label](url) → just the label (PDF hyperlinks need extra work)
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
        # Escape bare & that aren't already entities
        text = re.sub(r"&(?!(?:amp|lt|gt|bull|nbsp);)", "&amp;", text)
        return text

    def parse_table(lines: list[str]):
        """Parse a markdown table block into a ReportLab Table."""
        rows = []
        for line in lines:
            if re.match(r"^\|[\s\-:|]+\|$", line):
                continue  # separator row
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            rows.append(cells)
        if not rows:
            return None

        # Normalize column count
        col_count = max(len(r) for r in rows)
        for r in rows:
            while len(r) < col_count:
                r.append("")

        # Convert header cells to bold paragraphs, body cells to normal
        table_data = []
        for i, row in enumerate(rows):
            style = h3 if i == 0 else body
            table_data.append([Paragraph(md_inline(cell), style) for cell in row])

        available_width = letter[0] - 120  # account for margins
        col_width = available_width / col_count
        tbl = Table(table_data, colWidths=[col_width] * col_count)
        tbl.setStyle(TABLE_STYLE)
        return tbl

    story = []
    lines = request.content.split("\n")
    for start_i, line in enumerate(lines):
        if re.match(r"^#{1,3}\s+", line.strip()):
            lines = lines[start_i:]
            break
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Skip horizontal rules ---
        if re.match(r"^[-*_]{3,}$", stripped):
            story.append(Spacer(1, 8))
            i += 1
            continue

        # --- Fenced code blocks ```...``` ---
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing ```
            code_text = "\n".join(code_lines)
            # Escape XML special chars inside code blocks
            code_text = code_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(code_text.replace("\n", "<br/>"), code_style))
            story.append(Spacer(1, 6))
            continue

        # --- Markdown tables ---
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            tbl = parse_table(table_lines)
            if tbl:
                story.append(tbl)
                story.append(Spacer(1, 10))
            continue

        # --- Headings ---
        h_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if h_match:
            level = len(h_match.group(1))
            text = md_inline(h_match.group(2))
            style = h1 if level == 1 else h2 if level == 2 else h3
            story.append(Paragraph(text, style))
            i += 1
            continue

        # --- Bullet points (-, *, +) with optional nesting ---
        bullet_match = re.match(r"^(\s*)([-*+])\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = md_inline(bullet_match.group(3))
            b_style = ParagraphStyle(
                "BulletIndent",
                parent=bullet_style,
                leftIndent=20 + indent * 10,
            )
            story.append(Paragraph(f"• {text}", b_style))
            i += 1
            continue

        # --- Numbered lists ---
        num_match = re.match(r"^\s*(\d+)[.)]\s+(.*)", stripped)
        if num_match:
            num = num_match.group(1)
            text = md_inline(num_match.group(2))
            story.append(Paragraph(f"{num}. {text}", bullet_style))
            i += 1
            continue

        # --- Blockquotes ---
        if stripped.startswith(">"):
            quote_text = md_inline(stripped.lstrip("> ").strip())
            quote_style = ParagraphStyle(
                "Quote", parent=body,
                leftIndent=24, borderPad=4,
                borderColor=colors.HexColor("#aaaaaa"),
                borderWidth=0,
                textColor=colors.HexColor("#555555"),
                fontName="Helvetica-Oblique",
            )
            story.append(Paragraph(f"❝ {quote_text}", quote_style))
            i += 1
            continue

        # --- Empty lines → small spacer ---
        if not stripped:
            story.append(Spacer(1, 6))
            i += 1
            continue

        # --- Normal paragraph ---
        story.append(Paragraph(md_inline(stripped), body))
        i += 1

    doc.build(story)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=document.pdf"},
    )